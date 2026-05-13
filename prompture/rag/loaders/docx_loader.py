"""Microsoft Word (.docx) document loader."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from ..documents import Document, DocumentLoader, LoaderSource

_DOCX_HINT = "DOCXLoader requires 'python-docx'. Install with: pip install 'prompture[rag-docx]'"


def _require_docx():
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as e:  # pragma: no cover - exercised via test mocks
        raise ImportError(_DOCX_HINT) from e
    return DocxDocument


class DOCXLoader(DocumentLoader):
    """Load a .docx file into a single text ``Document``."""

    supported_extensions = (".docx",)

    def load(self, source: LoaderSource, **options: Any) -> list[Document]:
        DocxDocument = _require_docx()
        if isinstance(source, bytes):
            doc = DocxDocument(BytesIO(source))
            source_label = "<bytes>"
        else:
            path = Path(source)
            doc = DocxDocument(str(path))
            source_label = str(path)

        paragraphs = [p.text for p in doc.paragraphs if p.text]
        content = "\n\n".join(paragraphs)
        return [
            Document(
                content=content,
                metadata={
                    "source": source_label,
                    "paragraph_count": len(paragraphs),
                    "type": "docx",
                },
            )
        ]
